import os
import fitz  # PyMuPDF
from docx import Document # python-docx

def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text from PDF using PyMuPDF."""
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text("text")
        return text
    except Exception as e:
        print(f"❌ PDF Extraction Error: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from Word documents using python-docx."""
    try:
        doc = Document(file_path)
        # Extract text from paragraphs
        full_text = [para.text for para in doc.paragraphs]
        # Also extract text from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"❌ DOCX Extraction Error: {e}")
        return ""

def parse_resume(file_path: str) -> str:
    """
    Main entry point for Step 3. 
    Detects extension and routes to the correct extractor.
    """
    if not os.path.exists(file_path):
        print(f"❌ File not found at: {file_path}")
        return ""

    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        print(f"📄 Parsing PDF: {os.path.basename(file_path)}")
        return extract_text_from_pdf(file_path)
    
    elif ext == ".docx":
        print(f"📝 Parsing DOCX: {os.path.basename(file_path)}")
        return extract_text_from_docx(file_path)
    
    else:
        print(f"⚠️ Unsupported format: {ext}")
        return ""